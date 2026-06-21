from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ページ余白
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── ヘルパー ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_borders(cell, color='4472C4', sz='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_no_bottom_border(cell):
    """下枠線だけ消す（記入欄の区切りに使用）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '4472C4')
        tcBorders.append(b)
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'none')
    bot.set(qn('w:sz'), '0')
    bot.set(qn('w:space'), '0')
    bot.set(qn('w:color'), 'auto')
    tcBorders.append(bot)
    tcPr.append(tcBorders)

def write_label(cell, text, size=11):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'MS Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def write_value(cell, text='', size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'MS Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 914400 / 2.54 / 914400 * 914400 / 360)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def cm_to_twips(cm):
    return int(cm * 567)  # 1 cm = 567 twips (approx)

def set_row_height_twips(row, twips):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(twips))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

# ── タイトル ──────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('西小学校　読み聞かせ記録')
r.bold = True
r.font.size = Pt(18)
r.font.name = 'MS Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

doc.add_paragraph()

# ── メインフォームテーブル ─────────────────────────────────
# 列構成: [ラベル列, 値列, ラベル列, 値列]
# 列幅: 3cm, 5.5cm, 3cm, 5.5cm  = 17cm
tbl = doc.add_table(rows=0, cols=4)
tbl.style = 'Table Grid'

COL_W = [Cm(3), Cm(5.5), Cm(3), Cm(5.5)]
LABEL_BG = '4472C4'
VALUE_BG  = 'EEF3FB'

def add_row_2col(label1, label2, height_twips=700):
    """2カラム対称行を追加"""
    row = tbl.add_row()
    set_row_height_twips(row, height_twips)
    cells = row.cells
    for i, w in enumerate(COL_W):
        cells[i].width = w
    set_cell_bg(cells[0], LABEL_BG); set_borders(cells[0])
    set_cell_bg(cells[1], VALUE_BG); set_borders(cells[1])
    set_cell_bg(cells[2], LABEL_BG); set_borders(cells[2])
    set_cell_bg(cells[3], VALUE_BG); set_borders(cells[3])
    write_label(cells[0], label1)
    write_value(cells[1])
    write_label(cells[2], label2)
    write_value(cells[3])
    return row

def add_row_full(label, value_text='', height_twips=700):
    """ラベル1列 + 値3列結合の全幅行"""
    row = tbl.add_row()
    set_row_height_twips(row, height_twips)
    cells = row.cells
    for i, w in enumerate(COL_W):
        cells[i].width = w
    set_cell_bg(cells[0], LABEL_BG); set_borders(cells[0])
    # 右3セルを結合
    merged = cells[1].merge(cells[2]).merge(cells[3])
    set_cell_bg(merged, VALUE_BG); set_borders(merged)
    write_label(cells[0], label)
    write_value(merged, value_text)
    return row

def add_row_wide(label, height_twips=1800):
    """ラベル1列 + 値3列結合・高さ広め（記述欄）"""
    return add_row_full(label, height_twips=height_twips)

# 行1: 読み手名 ／ 日　時
add_row_2col('読み手名', '日　　時')

# 行2: 学　年 ／ 本　名
add_row_2col('学　　年', '本　　名')

# 行3: 貸出元（全幅）
row_kashidashi = tbl.add_row()
set_row_height_twips(row_kashidashi, 750)
cells = row_kashidashi.cells
for i, w in enumerate(COL_W):
    cells[i].width = w
set_cell_bg(cells[0], LABEL_BG); set_borders(cells[0])
merged_k = cells[1].merge(cells[2]).merge(cells[3])
set_cell_bg(merged_k, VALUE_BG); set_borders(merged_k)
write_label(cells[0], '貸　出　元')
# チェックボックス風テキスト
p = merged_k.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_k = p.add_run('  □ 自前　　□ 学校　　□ 絵本館　　□ 他（　　　　　　　　）')
run_k.font.size = Pt(11)
run_k.font.name = 'MS Gothic'
run_k._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

# 行4: 子どもたちの反応（大きめ記述欄）
add_row_wide('子どもたちの反応', height_twips=2400)

# 行5: その他
add_row_wide('そ　の　他', height_twips=1800)

# ── フッター ──────────────────────────────────────────────
doc.add_paragraph()
foot_p = doc.add_paragraph()
foot_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = foot_p.add_run('記録日：　　　　年　　月　　日')
fr.font.size = Pt(10)
fr.font.name = 'MS Gothic'
fr._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

doc.save('/home/user/word-scramble/西小学校読み聞かせ記録.docx')
print('Done')
