from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ余白を狭く
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# タイトル
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('読み聞かせ担当　日程表')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'MS Gothic'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

doc.add_paragraph()  # 空行

# ── ヘルパー: セルに色を付ける ──────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'MS Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

# ── 【学年順】表 ─────────────────────────────────────────
h1 = doc.add_paragraph()
h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = h1.add_run('【学年順】')
r1.bold = True
r1.font.size = Pt(12)
r1.font.name = 'MS Gothic'
r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

grade_data = [
    ('1年生', '波方',    '6月25日'),
    ('2年生', '滝上',    '6月24日'),
    ('3年生', '岡田',    '6月24日'),
    ('4年生', 'ロジャース', '7月8日'),
    ('5年生', '滝上',    '7月8日'),
    ('6年生', '赤前',    '6月23日'),
]

tbl1 = doc.add_table(rows=1 + len(grade_data), cols=3)
tbl1.style = 'Table Grid'

# 列幅
col_widths = [Cm(3), Cm(5), Cm(4)]
for i, w in enumerate(col_widths):
    for cell in tbl1.columns[i].cells:
        cell.width = w

# ヘッダー行
headers = ['学年', '担当者', '実施日']
header_cells = tbl1.rows[0].cells
for i, h in enumerate(headers):
    set_cell_bg(header_cells[i], '4472C4')
    set_cell_border(header_cells[i])
    cell_text(header_cells[i], h, bold=True, size=11, color='FFFFFF')

# データ行
row_colors = ['FFFFFF', 'DCE6F1']
for ri, (grade, name, date) in enumerate(grade_data):
    row = tbl1.rows[ri + 1]
    bg = row_colors[ri % 2]
    values = [grade, name, date]
    for ci, val in enumerate(values):
        set_cell_bg(row.cells[ci], bg)
        set_cell_border(row.cells[ci])
        cell_text(row.cells[ci], val, size=11)

doc.add_paragraph()  # 空行

# ── 【日付順】表 ─────────────────────────────────────────
h2 = doc.add_paragraph()
h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r2 = h2.add_run('【日付順】')
r2.bold = True
r2.font.size = Pt(12)
r2.font.name = 'MS Gothic'
r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

date_data = [
    ('6月23日', '6年生', '赤前'),
    ('6月24日', '2年生', '滝上'),
    ('6月24日', '3年生', '岡田'),
    ('6月25日', '1年生', '波方'),
    ('7月8日',  '4年生', 'ロジャース'),
    ('7月8日',  '5年生', '滝上'),
]

tbl2 = doc.add_table(rows=1 + len(date_data), cols=3)
tbl2.style = 'Table Grid'

for i, w in enumerate(col_widths):
    for cell in tbl2.columns[i].cells:
        cell.width = w

headers2 = ['実施日', '学年', '担当者']
header_cells2 = tbl2.rows[0].cells
for i, h in enumerate(headers2):
    set_cell_bg(header_cells2[i], '4472C4')
    set_cell_border(header_cells2[i])
    cell_text(header_cells2[i], h, bold=True, size=11, color='FFFFFF')

for ri, (date, grade, name) in enumerate(date_data):
    row = tbl2.rows[ri + 1]
    bg = row_colors[ri % 2]
    values = [date, grade, name]
    for ci, val in enumerate(values):
        set_cell_bg(row.cells[ci], bg)
        set_cell_border(row.cells[ci])
        cell_text(row.cells[ci], val, size=11)

# 保存
doc.save('/home/user/word-scramble/読み聞かせ日程表.docx')
print('Done')
